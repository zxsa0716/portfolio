# -*- coding: utf-8 -*-
"""
블로그 원고 마크다운 → Word(.docx) 변환기  (2종 동시 생성)

1) NN_슬러그.docx          ← 티스토리 붙여넣기용. 본문을 **마크다운 원문 그대로** 담는다.
                             제목(#/##/###), 굵게(**), 기울임(*), 밑줄(<u>), 구분선(---),
                             표(|...|), 인용(>), 목록(-) 모두 마크다운 문법을 문자로 유지.
                             글자 서식을 일절 입히지 않아 붙여넣을 때 문법이 깨지지 않는다.
                             이미지는 캡션 위치에 실제로 삽입 → 붙여넣으면 함께 업로드된다.

2) NN_슬러그_미리보기.docx  ← 사람이 읽고 검토하는 용도. 서식을 적용해 보기 좋게 조판.

사용: python md2docx.py
"""
import os, re, glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

BLOG = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FONT = "맑은 고딕"
GREEN = RGBColor(0x1F, 0x4E, 0x36)
GRAY = RGBColor(0x6E, 0x6E, 0x6E)
DARK = RGBColor(0x26, 0x26, 0x26)
MAXW_CM = 15.5


# ────────────────────────────── 공통 유틸 ──────────────────────────────
def _font(run, size, bold=False, italic=False, color=DARK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rf.set(qn(a), FONT)


def _para(doc, before=0, after=6, line=1.6, indent=0.0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = line
    if indent: pf.first_line_indent = Cm(indent)
    if align is not None: p.alignment = align
    return p


def _inline(p, text, size=11, base_bold=False):
    """**bold** 인라인 처리 (미리보기용)"""
    for seg in re.split(r'(\*\*[^*]+\*\*)', text):
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            _font(p.add_run(seg[2:-2]), size, bold=True)
        else:
            _font(p.add_run(seg), size, bold=base_bold)


def _botborder(p, color="1F4E36", sz=6):
    pPr = p._p.get_or_add_pPr(); b = OxmlElement('w:pBdr'); bt = OxmlElement('w:bottom')
    bt.set(qn('w:val'), 'single'); bt.set(qn('w:sz'), str(sz))
    bt.set(qn('w:space'), '3'); bt.set(qn('w:color'), color)
    b.append(bt); pPr.append(b)


def _shade(p, fill="F4F7F4"):
    pPr = p._p.get_or_add_pPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill); pPr.append(sh)


def setup(doc, narrow=False):
    s = doc.sections[0]
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    m = Cm(1.8) if narrow else Cm(2.4)
    s.left_margin = m; s.right_margin = m


def _img_width(path, scale=1.0):
    try:
        w, h = Image.open(path).size
        width = MAXW_CM
        if h / max(w, 1) > 2.2:          # 지나치게 긴 세로 이미지는 폭을 줄임
            width = MAXW_CM * 0.62
    except Exception:
        width = MAXW_CM
    return Cm(width * scale)


def _blocks(md_path):
    """post.md를 블록 목록으로 파싱. front matter는 버린다."""
    lines = open(md_path, encoding='utf-8').read().split('\n')
    out, i, in_front = [], 0, False
    while i < len(lines):
        ln = lines[i].rstrip()
        if i == 0 and ln.strip() == '---':
            in_front = True; i += 1; continue
        if in_front:
            if ln.strip() == '---': in_front = False
            i += 1; continue

        if not ln.strip():
            out.append(('blank', '')); i += 1; continue

        m = re.match(r'!\[(.*?)\]\((.*?)\)', ln.strip())
        if m:
            out.append(('image', (m.group(1), m.group(2)))); i += 1; continue

        # 표: 둘째 줄이 |---|---| 구분행
        if ln.strip().startswith('|') and i + 1 < len(lines) \
           and re.match(r'^\|[\s:\-|]+\|$', lines[i + 1].strip()):
            rows = [lines[i].rstrip(), lines[i + 1].rstrip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(lines[i].rstrip()); i += 1
            out.append(('table', rows)); continue

        out.append(('line', ln)); i += 1
    return out


# ─────────────────── 1) 티스토리 붙여넣기용: 마크다운 원문 ───────────────────
def _md_para(doc, text):
    """마크다운 한 줄을 서식 없는 순수 텍스트로. 붙여넣을 때 문법이 보존된다."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.35
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if text:
        _font(p.add_run(text), 10.5)      # 굵게·색·기울임 일절 없음
    return p


def build_markdown_docx(md_path, out_path):
    """본문 전체를 마크다운 문법 그대로 담고, 그림은 캡션 자리에 실제 삽입."""
    base = os.path.dirname(md_path)
    doc = Document(); setup(doc, narrow=True)
    fignum = 0

    for kind, val in _blocks(md_path):
        if kind == 'blank':
            _md_para(doc, ''); continue

        if kind == 'line':
            _md_para(doc, val); continue

        if kind == 'table':
            for row in val:
                _md_para(doc, row)
            continue

        # 이미지: 실제 그림 + 그 아래 마크다운 기울임 캡션
        caption, rel = val
        fignum += 1
        path = os.path.join(base, rel)
        if os.path.exists(path):
            ip = doc.add_paragraph()
            ipf = ip.paragraph_format
            ipf.space_before = Pt(0); ipf.space_after = Pt(0)
            ipf.line_spacing_rule = WD_LINE_SPACING.SINGLE
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.add_run().add_picture(path, width=_img_width(path))
        else:
            _md_para(doc, f'[이미지 없음: {rel}]')
        _md_para(doc, '')
        _md_para(doc, f'*그림 {fignum}. {caption}*')

    doc.save(out_path)
    return fignum


# ─────────────────── 2) 검토용 미리보기: 서식 적용 ───────────────────
def _prev_table(doc, rows):
    def cells(s): return [c.strip() for c in s.strip().strip('|').split('|')]
    data = [cells(rows[0])] + [cells(r) for r in rows[2:]]
    ncol = max(len(r) for r in data)
    t = doc.add_table(rows=len(data), cols=ncol)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(data):
        trPr = t.rows[ri]._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        if ri == 0: trPr.append(OxmlElement('w:tblHeader'))
        for ci in range(ncol):
            cell = t.cell(ri, ci); cell.text = ''
            p = cell.paragraphs[0]; pf = p.paragraph_format
            pf.space_before = Pt(2); pf.space_after = Pt(2)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            txt = row[ci] if ci < len(row) else ''
            if ri == 0:
                _shade(p, "EDF2ED"); _font(p.add_run(re.sub(r'\*\*', '', txt)), 10, bold=True, color=GREEN)
            else:
                _inline(p, txt, 10)
    _para(doc, before=0, after=10)


def build_preview_docx(md_path, out_path):
    base = os.path.dirname(md_path)
    doc = Document(); setup(doc)
    fignum = 0
    blocks = _blocks(md_path)
    k = 0
    while k < len(blocks):
        kind, val = blocks[k]; k += 1
        if kind == 'blank':
            continue

        if kind == 'table':
            _prev_table(doc, val); continue

        if kind == 'image':
            caption, rel = val
            fignum += 1
            path = os.path.join(base, rel)
            if not os.path.exists(path):
                p = _para(doc, after=4)
                _font(p.add_run(f"[이미지 없음: {rel}]"), 9.5, color=RGBColor(0xC0, 0x39, 0x39))
                continue
            p = _para(doc, before=8, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
            p.add_run().add_picture(path, width=_img_width(path))
            cp = _para(doc, before=0, after=12, line=1.35, align=WD_ALIGN_PARAGRAPH.CENTER)
            _font(cp.add_run(f"그림 {fignum}. "), 9.5, bold=True, color=GREEN)
            _font(cp.add_run(caption), 9.5, color=GRAY)
            continue

        ln = val
        if ln.startswith('#### '):
            p = _para(doc, before=10, after=3); _font(p.add_run(ln[5:].strip()), 11.5, bold=True, color=DARK); continue
        if ln.startswith('### '):
            p = _para(doc, before=12, after=4); _font(p.add_run(ln[4:].strip()), 12.5, bold=True, color=GREEN); continue
        if ln.startswith('## '):
            p = _para(doc, before=18, after=6)
            _font(p.add_run(ln[3:].strip()), 14, bold=True, color=GREEN); _botborder(p); continue
        if ln.startswith('# '):
            p = _para(doc, before=0, after=4); _font(p.add_run(ln[2:].strip()), 18, bold=True, color=GREEN); continue
        if re.match(r'^-{3,}$', ln.strip()):
            p = _para(doc, before=6, after=6); _botborder(p, "CCCCCC", 4); continue
        if ln.startswith('> '):
            buf = [ln[2:].strip()]
            while k < len(blocks) and blocks[k][0] == 'line' and blocks[k][1].startswith('> '):
                buf.append(blocks[k][1][2:].strip()); k += 1
            p = _para(doc, before=8, after=10, line=1.5); _shade(p); _inline(p, ' '.join(buf), 10.5); continue
        if re.match(r'^[-*] ', ln.strip()):
            p = _para(doc, before=0, after=4, line=1.5)
            p.paragraph_format.left_indent = Cm(0.6)
            _font(p.add_run('· '), 11, bold=True, color=GREEN); _inline(p, ln.strip()[2:], 11); continue
        if re.match(r'^\d+\. ', ln.strip()):
            p = _para(doc, before=0, after=4, line=1.5)
            p.paragraph_format.left_indent = Cm(0.6); _inline(p, ln.strip(), 11); continue

        p = _para(doc, before=0, after=9, line=1.65, indent=0.3, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _inline(p, ln.strip(), 11)

    doc.save(out_path)
    return fignum


def write_upload_order(md_path, out_path):
    """그림이 본문에 나오는 순서대로 파일명을 적어 둔다. 이미지가 붙여넣기로
    옮겨지지 않을 때 이 순서대로 직접 올리면 된다."""
    imgs = [v for k, v in _blocks(md_path) if k == 'image']
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write('# 그림 업로드 순서 (본문에 나오는 순서)\n')
        f.write('# 파일명 번호 = 본문 등장 순서입니다. 붙여넣기로 이미지가 함께\n')
        f.write('# 옮겨지지 않으면 01부터 차례대로 해당 캡션 위에 올리면 됩니다.\n\n')
        for i, (caption, rel) in enumerate(imgs, 1):
            head = re.sub(r'\s+', ' ', caption)[:52]
            f.write(f'그림 {i:2d}  {os.path.basename(rel):<10}  {head}…\n')
    return len(imgs)


def main():
    targets = sorted(glob.glob(os.path.join(BLOG, '*', 'post.md')))
    if not targets:
        print('post.md 없음'); return
    for t in targets:
        folder = os.path.basename(os.path.dirname(t))
        d = os.path.dirname(t)
        try:
            n = build_markdown_docx(t, os.path.join(d, f'{folder}.docx'))
            build_preview_docx(t, os.path.join(d, f'{folder}_미리보기.docx'))
            write_upload_order(t, os.path.join(d, '그림_업로드순서.txt'))
            chars = len(re.sub(r'\s', '', open(t, encoding='utf-8').read()))
            print(f'OK  {folder:24s} 그림 {n:2d}개 · 본문 {chars:,}자 → {folder}.docx + _미리보기.docx')
        except Exception as e:
            print(f'ERR {folder}: {e}')


if __name__ == '__main__':
    main()
